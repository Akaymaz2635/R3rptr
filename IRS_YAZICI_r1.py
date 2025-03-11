import customtkinter as ctk
from docx import Document
import docx
from tkinter import filedialog, ttk, Tk, messagebox
import tkinter as tk
import shutil
import os
import sys
import subprocess
import datetime
import logging


# Set up logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler("app.log"),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger("IRS_YAZICI")

# Set theme to modern look
ctk.set_appearance_mode("dark")
ctk.set_default_color_theme("blue")


class LotDetailManager:

    """
    A class to manage lot details separately from the main application.
    This creates better separation of concerns and makes the code more maintainable.
    """
    
    def __init__(self, parent):
        """
        Initialize the lot detail manager.
        
        Args:
            parent: The parent application that owns this manager
        """
        self.parent = parent
        self.lot_details = {}
        self.part_quantities = {}
        self.part_numbers = {}
        self.lot_notes = {}
        
    def show_lot_detail_dialog(self, row_idx, item_no, dimension, actual_value, widgets_dict):
        """
        Show the lot detail dialog for a specific item.
        
        Args:
            row_idx (int): The row index in the table
            dimension (str): The dimension value
            item_no (str): The item number
            actual_value (str): The current actual value
            widgets_dict (dict): Dictionary of widgets to update
        """
        # Ensure we have valid values for the key parts
        dimension = dimension if dimension else "Unknown Dimension"
        item_no = item_no if item_no else "Unknown Item"
        key = f"{dimension}_{item_no}"
        
        # Logging for debugging
        logger.debug(f"Opening lot detail dialog for key: {key}")
        logger.debug(f"Existing part quantities: {list(self.part_quantities.keys())}")
        
        # Create the dialog window
        detail_window = ctk.CTkToplevel(self.parent)
        detail_window.title(f"Lot Detail - {item_no}")
        detail_window.geometry("500x500")
        detail_window.grab_set()  # Make it modal
        
        # Create the main content frame
        content_frame = ctk.CTkFrame(detail_window, corner_radius=10)
        content_frame.pack(padx=20, pady=20, fill="both", expand=True)
        
        # Add title
        ctk.CTkLabel(content_frame, text=f"Lot Detail for Item: {item_no}", 
                    font=("Helvetica", 18, "bold")).pack(pady=(10, 20))
        
        # Create scrollable frame for content
        main_info_frame = ctk.CTkFrame(content_frame, corner_radius=8)
        main_info_frame.pack(padx=15, pady=10, fill="both", expand=True)
        
        info_scroll = ctk.CTkScrollableFrame(main_info_frame, corner_radius=8)
        info_scroll.pack(fill="both", expand=True)
        
        info_frame = ctk.CTkFrame(info_scroll, corner_radius=0)
        info_frame.pack(fill="both", expand=True)
        
        # Add item information
        ctk.CTkLabel(info_frame, text=f"DIMENSION: {dimension}", 
                    font=("Helvetica", 14)).pack(anchor="w", padx=15, pady=5)
        ctk.CTkLabel(info_frame, text=f"ITEM NO: {item_no}", 
                    font=("Helvetica", 14)).pack(anchor="w", padx=15, pady=5)
        ctk.CTkLabel(info_frame, text=f"ACTUAL Value: {actual_value}", 
                    font=("Helvetica", 14)).pack(anchor="w", padx=15, pady=5)
        ctk.CTkLabel(info_frame, text=f"Row Index: {row_idx}", 
                    font=("Helvetica", 14)).pack(anchor="w", padx=15, pady=5)
        
        # Part quantity section
        quantity_frame = ctk.CTkFrame(info_frame)
        quantity_frame.pack(fill="x", padx=15, pady=5)
        
        ctk.CTkLabel(quantity_frame, text="Part Quantity:", 
                    font=("Helvetica", 14)).pack(side="left", padx=(0, 10))
        
        # Get saved quantity if it exists
        saved_quantity = self.part_quantities.get(key, "0")
        quantity_var = tk.StringVar(value=saved_quantity)
        
        quantity_display = ctk.CTkLabel(quantity_frame, 
                                    textvariable=quantity_var,
                                    font=("Helvetica", 14, "bold"))
        quantity_display.pack(side="left", padx=10)
        
        # Parts list frame
        parts_scroll = ctk.CTkScrollableFrame(info_frame, corner_radius=6, height=150)
        parts_scroll.pack(fill="x", padx=15, pady=10)
        
        parts_frame = ctk.CTkFrame(parts_scroll, corner_radius=0)
        parts_frame.pack(fill="both", expand=True)
        
        # Get saved part numbers
        saved_part_numbers = self.part_numbers.get(key, {})
        part_entries = {}  # Will store the entry widgets
        
        # Function to update the part entries based on quantity
        def update_part_entries():
            # Clear existing entries
            for widget in parts_frame.winfo_children():
                widget.destroy()
                
            quantity = int(quantity_var.get())
            
            if quantity > 0:
                ctk.CTkLabel(parts_frame, text="Part Numbers:", 
                            font=("Helvetica", 14, "bold")).pack(anchor="w", pady=(0, 5))
            
            for i in range(1, quantity + 1):
                entry_frame = ctk.CTkFrame(parts_frame)
                entry_frame.pack(fill="x", pady=2)
                
                ctk.CTkLabel(entry_frame, text=f"{i}-", 
                            font=("Helvetica", 12)).pack(side="left", padx=(0, 5))
                
                entry = ctk.CTkEntry(entry_frame, width=200, height=30, corner_radius=6, 
                                    font=("Helvetica", 12))
                entry.pack(side="left", fill="x", expand=True)
                
                # Load saved value if available
                if str(i) in saved_part_numbers:
                    entry.insert(0, saved_part_numbers[str(i)])
                
                part_entries[i] = entry
        
        # Create increase/decrease buttons for quantity
        def increase_quantity():
            current = int(quantity_var.get())
            quantity_var.set(str(current + 1))
            update_part_entries()
            
        increase_btn = ctk.CTkButton(quantity_frame, text="Increase", 
                                    font=("Helvetica", 12),
                                    width=80, height=25, corner_radius=6,
                                    command=increase_quantity)
        increase_btn.pack(side="left", padx=5)
        
        def decrease_quantity():
            current = int(quantity_var.get())
            if current > 0:  # Prevent negative values
                quantity_var.set(str(current - 1))
                update_part_entries()
        
        decrease_btn = ctk.CTkButton(quantity_frame, text="Decrease", 
                                    font=("Helvetica", 12),
                                    width=80, height=25, corner_radius=6,
                                    command=decrease_quantity)
        decrease_btn.pack(side="left", padx=5)
        
        # Initialize part entries
        update_part_entries()
        
        # Notes section
        ctk.CTkLabel(info_frame, text="Additional Notes:", 
                    font=("Helvetica", 14)).pack(anchor="w", padx=15, pady=(15, 5))
        
        notes_frame = ctk.CTkFrame(info_frame)
        notes_frame.pack(padx=15, pady=5, fill="x")
        
        notes_entry = ctk.CTkTextbox(notes_frame, height=100, corner_radius=6)
        notes_entry.pack(side="left", fill="both", expand=True)
        
        # Load saved notes if available
        saved_notes = self.lot_notes.get(key, "")
        if saved_notes:
            notes_entry.insert("1.0", saved_notes)
        
        scrollbar = ctk.CTkScrollbar(notes_frame, command=notes_entry.yview)
        scrollbar.pack(side="right", fill="y")
        notes_entry.configure(yscrollcommand=scrollbar.set)
        
        # Function to save the data
        def save_lot_data():
            # Get notes
            notes_text = notes_entry.get("1.0", "end").strip()
            
            # Get part numbers
            current_part_numbers = {}
            for i, entry in part_entries.items():
                current_part_numbers[str(i)] = entry.get()
            
            # Calculate min and max values for ACTUAL
            valid_numbers = self._extract_numeric_values(current_part_numbers)
            
            if valid_numbers:
                min_value = min(valid_numbers)
                max_value = max(valid_numbers)
                min_max_str = f"{min_value} / {max_value}"
                
                # Update widgets and values
                if key in widgets_dict:
                    widgets_dict[key].delete(0, "end")
                    widgets_dict[key].insert(0, min_max_str)
                
                # Update the parent's actual_values dictionary
                self.parent.actual_values[key] = min_max_str
            
            # Save all lot details
            self.part_quantities[key] = quantity_var.get()
            self.part_numbers[key] = current_part_numbers
            self.lot_notes[key] = notes_text
            
            logger.info(f"Saved lot details for {key}")
            
            # Close the dialog
            detail_window.destroy()
        
        # Add save and close buttons
        button_frame = ctk.CTkFrame(content_frame)
        button_frame.pack(pady=15)
        
        ctk.CTkButton(button_frame, text="Save", font=("Helvetica", 14), 
                     width=100, height=30, corner_radius=6, 
                     command=save_lot_data).pack(side="left", padx=10)
        
        ctk.CTkButton(button_frame, text="Close", font=("Helvetica", 14), 
                     width=100, height=30, corner_radius=6, 
                     command=detail_window.destroy).pack(side="left", padx=10)
    
    def _extract_numeric_values(self, part_numbers):
        """
        Extract numeric values from part numbers, handling various formats.
        
        Args:
            part_numbers (dict): Dictionary of part numbers
            
        Returns:
            list: List of valid numeric values
        """
        valid_numbers = []
        
        for num_str in part_numbers.values():
            if not num_str or not num_str.strip():
                continue
                
            # Handle slash-separated values
            if '/' in num_str:
                parts = num_str.split('/')
                for part in parts:
                    try:
                        if part and part.strip():
                            valid_numbers.append(float(part.strip()))
                    except ValueError:
                        continue
            else:
                # Handle single values
                try:
                    valid_numbers.append(float(num_str))
                except ValueError:
                    continue
        
        return valid_numbers
    
    def export_lot_details_to_text(self, folder_path, project_info):
        """
        Export lot details to a text file using a tabular format.
        Overwrites any existing file with current data only.
        
        Args:
            folder_path (str): Path to save the file
            project_info (dict): Project information for the header
            
        Returns:
            str: Path to the created file
            list: Content lines of the lot details
        """
        lot_details_file = os.path.join(folder_path, "lot_details.txt")
        lot_details_content = []
        
        # Always create header with project info regardless of file existence
        lot_details_content.append("=" * 80)
        lot_details_content.append("PROJE BILGILERI".center(80))
        lot_details_content.append("=" * 80)
        lot_details_content.append("")
        
        # Format project info in a tabular way
        for key, value in project_info.items():
            lot_details_content.append(f"{key:20}: {value}")
        
        lot_details_content.append("")
        lot_details_content.append("=" * 80)
        lot_details_content.append("LOT DETAYLARI".center(80))
        lot_details_content.append("=" * 80)
        lot_details_content.append("")
        
        # Add current timestamp
        timestamp = datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        lot_details_content.append(f"OLCUM TARIHI: {timestamp}".center(80))
        lot_details_content.append("=" * 80)
        lot_details_content.append("")
        
        # Process each key and format in a compact way
        has_details = False
        
        for key in self.part_quantities.keys():
            has_details = True
            actual_value = self.parent.actual_values.get(key, "")
            quantity = self.part_quantities.get(key, "0")
            part_numbers = self.part_numbers.get(key, {})
            notes = self.lot_notes.get(key, "")
            
            # Add key as header
            lot_details_content.append(f"● {key} ●".center(80))
            lot_details_content.append("-" * 80)
            
            # Create a two-column layout
            lot_details_content.append(f"{'ACTUAL DEGERI:':<20} {actual_value:<30}{'PARCA MIKTARI:':<15} {quantity}")
            
            # Prepare part numbers in a more organized way
            if part_numbers:
                lot_details_content.append("")
                lot_details_content.append("PARCA NUMARALARI:".center(80))
                
                # Display part numbers in a more table-like format
                rows = []
                for i in range(1, int(quantity) + 1, 2):
                    if str(i) in part_numbers and i+1 <= int(quantity) and str(i+1) in part_numbers:
                        rows.append(f"{i:2}- {part_numbers[str(i)]:<35} {i+1:2}- {part_numbers[str(i+1)]}")
                    elif str(i) in part_numbers:
                        rows.append(f"{i:2}- {part_numbers[str(i)]}")
                
                for row in rows:
                    lot_details_content.append(row)
            
            # Add notes if present
            if notes:
                lot_details_content.append("")
                lot_details_content.append("NOTLAR:".center(80))
                # Wrap notes at 78 characters for better presentation
                for i in range(0, len(notes), 78):
                    lot_details_content.append(notes[i:i+78])
            
            lot_details_content.append("")
            lot_details_content.append("-" * 80)
            lot_details_content.append("")
        
        # Write to file (always overwrite)
        try:
            # Create new file (overwrite mode)
            with open(lot_details_file, "w", encoding="utf-8") as f:
                for line in lot_details_content:
                    f.write(f"{line}\n")
            
            if has_details:
                logger.info(f"Created new lot details file with {len(self.part_quantities)} entries: {lot_details_file}")
            else:
                logger.info(f"Created new empty lot details file: {lot_details_file}")
            
            return lot_details_file, lot_details_content
        except Exception as e:
            logger.error(f"Error saving lot details: {str(e)}")
            return None, lot_details_content

    def load_lot_details_from_text(self, folder_path):
        """
        Load lot details from a text file, handling multiple appended entries.
        
        Args:
            folder_path (str): Path to the folder containing the lot_details.txt file
            
        Returns:
            bool: True if successful, False otherwise
        """
        lot_details_file = os.path.join(folder_path, "lot_details.txt")
        
        # Check if file exists
        if not os.path.exists(lot_details_file):
            logger.info(f"Lot details file not found at: {lot_details_file}")
            return False
        
        try:
            # Read file content
            with open(lot_details_file, "r", encoding="utf-8") as f:
                content = f.readlines()
            
            # Parse content
            current_key = None
            in_part_numbers = False
            in_notes = False
            notes_content = []
            
            for line in content:
                line = line.strip()
                
                # Skip empty lines
                if not line:
                    in_part_numbers = False
                    in_notes = False
                    
                    # Save notes if we were collecting them
                    if notes_content and current_key:
                        self.lot_notes[current_key] = "\n".join(notes_content)
                        notes_content = []
                    continue
                
                # Skip separators but reset state
                if line.startswith("=") or line.startswith("-"):
                    in_part_numbers = False
                    in_notes = False
                    
                    # Save notes if we were collecting them
                    if notes_content and current_key:
                        self.lot_notes[current_key] = "\n".join(notes_content)
                        notes_content = []
                    continue
                
                # Skip headers and timestamps
                if line == "LOT DETAYLARI:" or line.startswith("Güncelleme Tarihi:"):
                    continue
                    
                # Skip project info lines
                if any(line.startswith(prefix) for prefix in [
                    "PROJE TIPI:", "PARCA NUMARASI:", "OPERASYON NO:", 
                    "SERI NO:", "OLCUM TARIHI:"
                ]):
                    continue
                
                # Check if this is a key line (e.g., "6.3 RA_KN998")
                if "_" in line and not line.startswith(" ") and not ":" in line:
                    # Save notes if we were collecting them for a previous key
                    if notes_content and current_key:
                        self.lot_notes[current_key] = "\n".join(notes_content)
                        notes_content = []
                    
                    # Set new current key
                    current_key = line
                    in_part_numbers = False
                    in_notes = False
                    continue
                
                # Skip if we don't have a current key
                if not current_key:
                    continue
                
                # Check for ACTUAL value
                if line.startswith("ACTUAL Değeri:"):
                    value = line.replace("ACTUAL Değeri:", "").strip()
                    self.parent.actual_values[current_key] = value
                    continue
                
                # Check for part quantity
                if line.startswith("Parça Miktarı:"):
                    quantity = line.replace("Parça Miktarı:", "").strip()
                    self.part_quantities[current_key] = quantity
                    continue
                
                # Check for part numbers section
                if line == "Parça Numaraları:":
                    in_part_numbers = True
                    in_notes = False
                    
                    # Initialize part numbers dictionary if needed
                    if current_key not in self.part_numbers:
                        self.part_numbers[current_key] = {}
                    continue
                
                # Parse part numbers
                if in_part_numbers:
                    # Part number format should be "1- 12345"
                    parts = line.split("- ", 1)
                    if len(parts) == 2 and parts[0].isdigit():
                        part_idx = parts[0]
                        part_num = parts[1].strip()
                        self.part_numbers[current_key][part_idx] = part_num
                    continue
                
                # Check for notes section
                if line == "Notlar:":
                    in_part_numbers = False
                    in_notes = True
                    notes_content = []  # Reset notes for this key
                    continue
                
                # Collect notes content
                if in_notes:
                    notes_content.append(line)
            
            # Save any remaining notes
            if notes_content and current_key:
                self.lot_notes[current_key] = "\n".join(notes_content)
            
            # Count loaded items
            num_actual_values = len([v for v in self.parent.actual_values.values() if v])
            num_part_quantities = len(self.part_quantities)
            
            logger.info(f"Successfully loaded lot details from: {lot_details_file}")
            logger.info(f"Loaded: {num_actual_values} ACTUAL values, {num_part_quantities} lot details")
            
            # Optionally show notification when details are loaded
            if self.part_quantities or self.part_numbers or any(self.parent.actual_values.values()):
                messagebox.showinfo(
                    "Bilgi", 
                    f"Önceden kaydedilmiş lot detayları yüklendi.\n"
                    f"ACTUAL Değerleri: {num_actual_values}\n"
                    f"Lot Detayları: {num_part_quantities}"
                )
                
            return True
            
        except Exception as e:
            logger.error(f"Error loading lot details: {str(e)}")
            return False

    # def add_lot_details_to_word(self, doc, lot_details_content):
    #     """
    #     Add lot details to a Word document in a more space-efficient tabular layout
        
    #     Args:
    #         doc: The Word document object
    #         lot_details_content: List of content lines  
                
    #     Returns:
    #         bool: True if successful, False otherwise
    #     """
    #     try:
    #         # Add page break before lot details
    #         doc.add_paragraph().runs[0].add_break(docx.enum.text.WD_BREAK.PAGE)
            
    #         # Process content by sections
    #         section_start = False
    #         in_project_info = True
    #         current_key = None
    #         current_section = []
            
    #         for line in lot_details_content:
    #             # Check for section start/end markers
    #             if line.startswith("====="):
    #                 if "LOT DETAYLARI" in line or line.strip() == "=" * 80:
    #                     # This is a major section separator
                        
    #                     # If we were collecting a section, add it now
    #                     if current_section:
    #                         if in_project_info:
    #                             self._add_project_info_to_word(doc, current_section)
    #                             in_project_info = False
    #                         else:
    #                             self._add_update_timestamp_to_word(doc, current_section)
                            
    #                         current_section = []
                        
    #                     # Start a new section if this is the LOT DETAYLARI header
    #                     if "LOT DETAYLARI" in line:
    #                         # Add LOT DETAYLARI header
    #                         p = doc.add_paragraph()
    #                         p.add_run("LOT DETAYLARI").bold = True
    #                         p.style = 'Heading 1'
    #                         p.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
                            
    #                         # Add a horizontal line
    #                         border = p.paragraph_format.border
    #                         border.bottom.style = docx.enum.text.WD_BORDER_STYLE.SINGLE
    #                         border.bottom.width = docx.shared.Pt(1)
                        
    #                     continue
                
    #             # Check for key markers (● KEY ●)
    #             if line.strip().startswith("●") and line.strip().endswith("●"):
    #                 # If we were collecting a section, add it now
    #                 if current_section and current_key:
    #                     self._add_lot_detail_section_to_word(doc, current_key, current_section)
                    
    #                 # Extract the new key
    #                 current_key = line.strip().replace("●", "").strip()
    #                 current_section = []
    #                 continue
                
    #             # Check for update timestamp
    #             if "Güncelleme Tarihi:" in line:
    #                 # Start collecting update timestamp section
    #                 current_section = [line]
    #                 in_project_info = False
    #                 current_key = None
    #                 continue
                
    #             # Collect project info lines
    #             if in_project_info and not line.startswith("="):
    #                 current_section.append(line)
    #                 continue
                
    #             # Collect lot detail section lines
    #             if current_key is not None:
    #                 current_section.append(line)
            
    #         # Add any remaining section
    #         if current_section:
    #             if current_key:
    #                 self._add_lot_detail_section_to_word(doc, current_key, current_section)
    #             elif "Güncelleme Tarihi:" in current_section[0]:
    #                 self._add_update_timestamp_to_word(doc, current_section)
            
    #         # Add page break at end
    #         doc.add_paragraph().runs[0].add_break(docx.enum.text.WD_BREAK.PAGE)
            
    #         logger.info("Added lot details to Word document")
    #         return True
        
    #     except Exception as e:
    #         logger.error(f"Error adding lot details to Word: {str(e)}")
    #         return False

    def _add_project_info_to_word(self, doc, section_lines):
        """
        Add project info section to Word document
        
        Args:
            doc: The Word document object
            section_lines: List of project info lines
        """
        # Add title
        p = doc.add_paragraph()
        p.add_run("PROJE BİLGİLERİ").bold = True
        p.style = 'Heading 2'
        p.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
        
        # Create a table for project info
        table = doc.add_table(rows=0, cols=2)
        table.style = 'Table Grid'
        table.autofit = True
        
        # Add project info rows
        for line in section_lines:
            if ":" in line and not line.startswith("LOT DETAYLARI"):
                key, value = line.split(":", 1)
                row = table.add_row().cells
                row[0].text = key.strip()
                row[0].paragraphs[0].runs[0].bold = True
                row[1].text = value.strip()
        
        # Add space after table
        doc.add_paragraph()    

    def _add_update_timestamp_to_word(self, doc, section_lines):
        """
        Add update timestamp to Word document
        
        Args:
            doc: The Word document object
            section_lines: List containing the timestamp line
        """
        if section_lines and "Güncelleme Tarihi:" in section_lines[0]:
            p = doc.add_paragraph()
            run = p.add_run(section_lines[0].strip())
            run.bold = True
            run.italic = True
            p.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
            
            # Add a light horizontal line
            border = p.paragraph_format.border
            border.bottom.style = docx.enum.text.WD_BORDER_STYLE.SINGLE
            border.bottom.width = docx.shared.Pt(0.5)
            
            # Add space after timestamp
            doc.add_paragraph()

    def _add_lot_detail_section_to_word(self, doc, key, section_lines):
        """
        Add a lot detail section to Word document in a tabular format
        
        Args:
            doc: The Word document object
            key: The lot detail key
            section_lines: List of section content lines
        """
        # Add key as heading
        p = doc.add_paragraph()
        run = p.add_run(key)
        run.bold = True
        run.font.size = docx.shared.Pt(14)
        p.style = 'Heading 3'
        
        # Create a table for this lot detail
        table = doc.add_table(rows=0, cols=4)
        table.style = 'Table Grid'
        
        # First row for ACTUAL and quantity
        actual_value = ""
        quantity = ""
        
        # Find the ACTUAL and quantity values
        for line in section_lines:
            if "ACTUAL Değeri:" in line and "Parça Miktarı:" in line:
                parts = line.split("Parça Miktarı:")
                actual_part = parts[0].strip()
                quantity_part = "Parça Miktarı:" + parts[1].strip()
                
                actual_value = actual_part.replace("ACTUAL Değeri:", "").strip()
                quantity = quantity_part.replace("Parça Miktarı:", "").strip()
                break
        
        # Add ACTUAL and quantity row
        row = table.add_row().cells
        row[0].text = "ACTUAL Değeri:"
        row[0].paragraphs[0].runs[0].bold = True
        row[1].text = actual_value
        row[2].text = "Parça Miktarı:"
        row[2].paragraphs[0].runs[0].bold = True
        row[3].text = quantity
        
        # Process part numbers if present
        in_part_numbers = False
        part_numbers = []
        
        for line in section_lines:
            if line.strip() == "Parça Numaraları:":
                in_part_numbers = True
                continue
            
            if in_part_numbers and line.strip() and "-" in line:
                part_numbers.append(line.strip())
        
        # Add part numbers in a two-column layout
        if part_numbers:
            # Add a Part Numbers row header
            row = table.add_row().cells
            row[0].merge(row[3])
            row[0].text = "Parça Numaraları:"
            row[0].paragraphs[0].runs[0].bold = True
            row[0].paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
            
            # Add part numbers in pairs
            for i in range(0, len(part_numbers), 2):
                row = table.add_row().cells
                row[0].merge(row[1])
                row[0].text = part_numbers[i]
                
                if i + 1 < len(part_numbers):
                    row[2].merge(row[3])
                    row[2].text = part_numbers[i + 1]
        
        # Find notes if present
        in_notes = False
        notes_content = []
        
        for line in section_lines:
            if line.strip() == "Notlar:":
                in_notes = True
                continue
            
            if in_notes and line.strip() and not line.startswith("-"):
                notes_content.append(line.strip())
        
        # Add notes if present
        if notes_content:
            row = table.add_row().cells
            row[0].merge(row[3])
            row[0].text = "Notlar:"
            row[0].paragraphs[0].runs[0].bold = True
            row[0].paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
            
            row = table.add_row().cells
            row[0].merge(row[3])
            row[0].text = "\n".join(notes_content)
        
        # Add space after the table
        doc.add_paragraph()

    def export_lot_details_to_pdf(self, folder_path, project_info, lot_details_content=None):
        """
        Export lot details to a PDF file.
        Creates a simple PDF that maintains the exact text format of the TXT file.
        
        Args:
            folder_path (str): Path to save the file
            project_info (dict): Project information for the header
            lot_details_content (list, optional): Content lines of the lot details
            
        Returns:
            str: Path to the created PDF file or None if failed
        """
        try:
            # Define the output file path
            pdf_file = os.path.join(folder_path, "lot_details.pdf")
            
            # If lot_details_content is not provided, load from text file
            if not lot_details_content:
                lot_details_file = os.path.join(folder_path, "lot_details.txt")
                if os.path.exists(lot_details_file):
                    with open(lot_details_file, "r", encoding="utf-8") as f:
                        lot_details_content = f.read().splitlines()
                else:
                    logger.error("Cannot find lot_details.txt file")
                    return None
            
            # Method 1: Using reportlab to create a monospaced text PDF
            try:
                from reportlab.lib.pagesizes import A4
                from reportlab.lib.units import cm
                from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
                from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
                from reportlab.lib.enums import TA_LEFT, TA_CENTER
                
                # Create a PDF document with larger margins for readability
                doc = SimpleDocTemplate(
                    pdf_file,
                    pagesize=A4,
                    leftMargin=2*cm,
                    rightMargin=2*cm,
                    topMargin=2*cm,
                    bottomMargin=2*cm
                )
                
                # Create a monospace style that preserves whitespace
                mono_style = ParagraphStyle(
                    'MonoStyle',
                    fontName='Courier',
                    fontSize=9,
                    leading=11,  # Line spacing
                    alignment=TA_LEFT,
                    spaceAfter=0,
                    spaceBefore=0,
                    leftIndent=0,
                    rightIndent=0,
                    firstLineIndent=0,
                )
                
                # Create a centered monospace style for headers
                mono_center_style = ParagraphStyle(
                    'MonoCenterStyle',
                    parent=mono_style,
                    alignment=TA_CENTER,
                    fontSize=9,
                    leading=11,
                )
                
                # Build content
                elements = []
                
                # Process each line, preserving exact formatting including whitespace
                for line in lot_details_content:
                    if not line.strip():
                        # Empty lines need a small spacer
                        elements.append(Spacer(1, 11))
                        continue
                        
                    # Check if line is centered (contains many spaces on both sides)
                    if line.startswith("     ") and line.endswith("     "):
                        # Use centered style for headers and separators
                        style = mono_center_style
                    else:
                        style = mono_style
                    
                    # Replace spaces with non-breaking spaces to preserve formatting
                    # and escape any special HTML characters
                    line = line.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                    
                    # Use a special trick to preserve multiple spaces in reportlab
                    line_with_spaces = ""
                    for char in line:
                        if char == " ":
                            line_with_spaces += "&nbsp;"
                        else:
                            line_with_spaces += char
                    
                    # Add the paragraph
                    para = Paragraph(line_with_spaces, style)
                    elements.append(para)
                
                # Build the PDF
                doc.build(elements)
                logger.info(f"Successfully created simple text PDF using reportlab: {pdf_file}")
                return pdf_file
                
            except ImportError:
                logger.error("ReportLab not available for PDF generation. Trying alternative method...")
            except Exception as e:
                logger.error(f"Error using reportlab to create PDF: {e}")
            
            # Method 2: Fallback to a very simple method - create a basic PDF with PyFPDF if available
            try:
                import fpdf
                
                # Create a new PDF
                pdf = fpdf.FPDF(format='A4')
                pdf.add_page()
                
                # Set font to a monospaced font
                pdf.set_font('Courier', '', 10)
                
                # Calculate line height
                line_height = 5
                
                # Add content
                if lot_details_content:
                    for line in lot_details_content:
                        # Make sure line is properly encoded for FPDF
                        try:
                            pdf.cell(0, line_height, line, ln=1)
                        except Exception:
                            # If encoding issues, try to handle the line differently
                            try:
                                pdf.cell(0, line_height, line.encode('latin-1', 'replace').decode('latin-1'), ln=1)
                            except:
                                # If all else fails, just add a placeholder
                                pdf.cell(0, line_height, "[Gösterilemeyen satır]", ln=1)
                
                # Save the PDF
                pdf.output(pdf_file)
                logger.info(f"Successfully created basic PDF using PyFPDF: {pdf_file}")
                return pdf_file
            except ImportError:
                logger.error("PyFPDF not available for PDF generation")
            except Exception as e:
                logger.error(f"Error using PyFPDF to create PDF: {e}")
            
            # All methods failed
            logger.error("All PDF generation methods failed")
            return None
            
        except Exception as e:
            logger.error(f"Error creating PDF: {e}")
            return None


class ProjectManager:
    """
    Manages project structure, folders, and files
    """
    def __init__(self):
        self.project_info = {}
        self.serial_folder = None
    
    def create_project_structure(self, project_type, part_number, operation_number, serial_number, continue_measurement):
        """
        Create the project folder structure and info file
        
        Args:
            project_type (str): Project type
            part_number (str): Part number
            operation_number (str): Operation number
            serial_number (str): Serial number
            continue_measurement (bool): Whether to continue measurement
            
        Returns:
            tuple: (success, folder_path, error_message)
        """
        try:
            # Validate inputs
            if not project_type or not part_number or not operation_number or not serial_number:
                return False, None, "All fields must be filled"
            
            # Store project info
            self.project_info = {
                "PROJE TIPI": project_type,
                "PARCA NUMARASI": part_number,
                "OPERASYON NO": operation_number,
                "SERI NO": serial_number,
                "OLUSTURMA TARIHI": datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')
            }
            
            # Get desktop path
            desktop_path = os.path.join(os.path.expanduser("~"), "Desktop")
            
            # Create folder structure
            report_folder = os.path.join(desktop_path, "Report")
            if not os.path.exists(report_folder):
                os.makedirs(report_folder)
            
            project_type_folder = os.path.join(report_folder, project_type)
            if not os.path.exists(project_type_folder):
                os.makedirs(project_type_folder)
            
            part_number_folder = os.path.join(project_type_folder, part_number)
            if not os.path.exists(part_number_folder):
                os.makedirs(part_number_folder)
            
            operation_folder = os.path.join(part_number_folder, operation_number)
            if not os.path.exists(operation_folder):
                os.makedirs(operation_folder)
            
            serial_folder = os.path.join(operation_folder, serial_number)
            if not os.path.exists(serial_folder):
                os.makedirs(serial_folder)
            
            # Store serial folder for future use
            self.serial_folder = serial_folder
            
            logger.info(f"Project structure created at: {serial_folder}")
            return True, serial_folder, None
            
        except Exception as e:
            logger.error(f"Error creating project structure: {str(e)}")
            return False, None, str(e)
    
    def initialize_from_existing(self, project_type, part_number, operation_number, serial_number):
        """
        Initialize project information from existing project and check for lot details
        
        Args:
            project_type (str): Project type
            part_number (str): Part number
            operation_number (str): Operation number
            serial_number (str): Serial number
            
        Returns:
            tuple: (folder_exists, folder_path)
        """
        try:
            # Store project info
            self.project_info = {
                "PROJE TIPI": project_type,
                "PARÇA NUMARASI": part_number,
                "OPERASYON NO": operation_number,
                "SERI NO": serial_number,
            }
            
            # Get desktop path
            desktop_path = os.path.join(os.path.expanduser("~"), "Desktop")
            
            # Check if folder structure exists
            report_folder = os.path.join(desktop_path, "Report")
            project_type_folder = os.path.join(report_folder, project_type)
            part_number_folder = os.path.join(project_type_folder, part_number)
            operation_folder = os.path.join(part_number_folder, operation_number)
            serial_folder = os.path.join(operation_folder, serial_number)
            
            if os.path.exists(serial_folder):
                # Store serial folder for future use
                self.serial_folder = serial_folder
                logger.info(f"Found existing project structure at: {serial_folder}")
                return True, serial_folder
            else:
                logger.info(f"Project structure does not exist at: {serial_folder}")
                return False, None
                
        except Exception as e:
            logger.error(f"Error checking project structure: {str(e)}")
            return False, None
    
    def open_folder(self, folder_path):
        """Open folder in file explorer"""
        try:
            if os.name == 'nt':  # Windows
                os.startfile(folder_path)
            elif os.name == 'posix':  # macOS, Linux
                if sys.platform == 'darwin':  # macOS
                    subprocess.call(['open', folder_path])
                else:  # Linux
                    subprocess.call(['xdg-open', folder_path])
            logger.info(f"Opened folder: {folder_path}")
            return True
        except Exception as e:
            logger.error(f"Error opening folder: {str(e)}")
            return False


class TableReader:
    """
    Reads and processes tables from Word documents
    """
    def __init__(self):
        self.table_data = []
    
    def read_word_tables(self, docx_path):
        """
        Read tables from a Word document
        
        Args:
            docx_path (str): Path to the Word document
            
        Returns:
            list: List of tables
        """
        try:
            doc = Document(docx_path)
            table_data = []
            
            for table in doc.tables:
                processed_table = self._process_table(table)
                if processed_table:  # Only add non-empty tables
                    table_data.append(processed_table)
            
            self.table_data = table_data
            logger.info(f"Read {len(table_data)} tables from {docx_path}")
            return table_data
        except Exception as e:
            logger.error(f"Error reading Word document: {str(e)}")
            raise
    
    def _process_table(self, table):
        """
        Process a single table, removing RECORDING SHEET rows and adjusting headers
        
        Args:
            table: Word table object
            
        Returns:
            list: Processed table data
        """
        table_data = []
        
        # Extract rows, skipping RECORDING SHEET
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            
            # Skip empty rows
            if not any(row_data):
                continue
            
            # Check if this is a RECORDING SHEET row
            row_text = " ".join(row_data).upper()
            if any(pattern in row_text for pattern in ["RECORDIG SHEET", "RECORDING SHEET", "REC ORDIG SHEET"]):
                continue
            
            # Add valid row
            table_data.append(row_data)
        
        # Skip empty tables
        if not table_data:
            return None
        
        # Check if headers need adjustment
        if len(table_data) >= 2:
            first_row = table_data[0]
            second_row = table_data[1]
            
            # Check if ACTUAL is in the rows
            has_actual_first = any("ACTUAL" in cell.upper() for cell in first_row if cell)
            has_actual_second = any("ACTUAL" in cell.upper() for cell in second_row if cell)
            
            # If first row doesn't have ACTUAL but second does, skip first row
            if not has_actual_first and has_actual_second:
                logger.info("First row doesn't have ACTUAL, using second row as header")
                table_data = table_data[1:]
        
        return table_data
    
    def get_column_indices(self, headers):
        """
        Find important column indices in the headers with improved B/P ZONE detection
        
        Args:
            headers (list): List of header strings
            
        Returns:
            dict: Dictionary of column indices
        """
        indices = {
            'actual': -1,
            'item_no': -1,
            'dimension': -1,
            'badge': -1,
            'tooling': -1,
            'insp_level': -1,
            'bp_zone': -1  # Added BP ZONE detection
        }
        
        for idx, header in enumerate(headers):
            header_upper = header.strip().upper()
            
            if "ACTUAL" in header_upper:
                indices['actual'] = idx
            if "ITEM NO" in header_upper or "ITEMNO" in header_upper:
                indices['item_no'] = idx
            if "DIMENSION" in header_upper or "DIM" in header_upper:
                indices['dimension'] = idx
            if "BADGE" in header_upper:
                indices['badge'] = idx
            if "TOOLING" in header_upper:
                indices['tooling'] = idx
            if "INSP. LEVEL" in header_upper or "INSPLEVEL" in header_upper:
                indices['insp_level'] = idx
            # Improved B/P ZONE detection
            if "B/P ZONE" in header_upper or "BP ZONE" in header_upper or header_upper == "B/P":
                indices['bp_zone'] = idx
                logger.info(f"Found B/P ZONE column at index {idx}: '{header}'")
        
        # As a fallback, look for any header containing "ZONE"
        if indices['bp_zone'] == -1:
            for idx, header in enumerate(headers):
                header_upper = header.strip().upper()
                if "ZONE" in header_upper:
                    indices['bp_zone'] = idx
                    logger.info(f"Found ZONE column at index {idx}: '{header}'")
                    break
        
        logger.info(f"Column indices: {indices}")
        return indices


class ReportGenerator:
    """
    Generates reports by updating Word documents with ACTUAL values
    """
    def __init__(self, project_manager):
        self.project_manager = project_manager
    
    def create_report(self, source_file, table_index, actual_values, col_indices, lot_detail_manager=None):
        """
        Create a report by updating a Word document with ALL tables.
        If file already exists in output folder, overwrite it instead of creating a copy.
        
        Args:
            source_file (str): Source Word document path
            table_index (int): Index of the currently selected table (used only for logging)
            actual_values (dict): Dictionary of actual values
            col_indices (dict): Dictionary of column indices
            lot_detail_manager (LotDetailManager): Lot detail manager for adding lot details
            
        Returns:
            tuple: (success, file_path, message)
        """
        try:
            # Validate serial folder
            if not self.project_manager.serial_folder:
                return False, None, "Project folder not created"
            
            # Find the serial number with case-insensitive matching to prevent key errors
            serial_number = None
            for key in self.project_manager.project_info:
                if key.upper() in ('SERI NO', 'SERIAL NO', 'SERINO'):
                    serial_number = self.project_manager.project_info[key]
                    break
            
            # If serial number is not found, use a default value
            if not serial_number:
                serial_number = "Report"
                logger.warning("Serial number not found in project info, using default")
            
            # Create new file name
            file_name = f"{serial_number}_{os.path.basename(source_file)}"
            output_file = os.path.join(self.project_manager.serial_folder, file_name)
            
            # Check if the output file already exists
            if os.path.exists(output_file):
                # If file exists, use it directly instead of copying
                logger.info(f"Output file already exists, overwriting: {output_file}")
                doc = Document(output_file)
            else:
                # If file doesn't exist, copy the source file
                logger.info(f"Creating new output file: {output_file}")
                shutil.copy2(source_file, output_file)
                doc = Document(output_file)
            
            # Initialize counters for statistics
            total_updated_count = 0
            updated_tables_count = 0
            
            # Process ALL tables in the document
            for current_table_idx, table in enumerate(doc.tables):
                # Find the header row for current table
                header_row_idx = self._find_header_row(table)
                if header_row_idx == -1:
                    logger.warning(f"Header row not found in table {current_table_idx}, skipping")
                    continue
                
                # Get column indices for current table
                current_col_indices = self._get_column_indices_for_table(table, header_row_idx)
                
                # Skip tables without required columns
                if current_col_indices['actual'] == -1 or current_col_indices['item_no'] == -1 or current_col_indices['dimension'] == -1:
                    logger.warning(f"Required columns not found in table {current_table_idx}, skipping")
                    continue
                
                # Update ACTUAL values in the current table
                updated_count = self._update_actual_values(
                    table, 
                    header_row_idx, 
                    actual_values, 
                    current_col_indices
                )
                
                if updated_count > 0:
                    total_updated_count += updated_count
                    updated_tables_count += 1
                    logger.info(f"Updated {updated_count} rows in table {current_table_idx}")
            
            # Add lot details if available
            lot_details_file = None
            lot_details_content = []
            lot_details_pdf = None
            
            if lot_detail_manager and hasattr(lot_detail_manager, 'part_quantities'):
                # Export lot details to text 
                lot_details_file, lot_details_content = lot_detail_manager.export_lot_details_to_text(
                    self.project_manager.serial_folder,
                    self.project_manager.project_info
                )
                
                # Export lot details to PDF
                lot_details_pdf = lot_detail_manager.export_lot_details_to_pdf(
                    self.project_manager.serial_folder,
                    self.project_manager.project_info,
                    lot_details_content
                )
                
                # If export was successful, read the entire file to get all lot details
#                if lot_details_file and lot_details_content:
                    # Add lot details to Word document
#                    lot_detail_manager.add_lot_details_to_word(doc, lot_details_content)
            
            # Add nonconforming items if available
            nonconforming_count = self.project_manager.project_info.get('UYGUNSUZ_OLCUM_SAYISI', 0)
            nonconforming_items = self.project_manager.project_info.get('UYGUNSUZ_OLCUMLER', '')
            
            # if nonconforming_items:
            #     # Add a page for nonconforming items
            #     doc.add_paragraph().runs[0].add_break(docx.enum.text.WD_BREAK.PAGE)
                
            #     heading = doc.add_paragraph()
            #     heading.add_run("UYGUNSUZ ÖLÇÜMLER").bold = True
            #     heading.style = 'Heading 1'
            #     heading.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
                
            #     doc.add_paragraph(f"Toplam uygunsuz ölçüm sayısı: {nonconforming_count}")
                
            #     for line in nonconforming_items.splitlines():
            #         if line.strip():
            #             doc.add_paragraph(line, style='List Bullet')
            
            # # Save the document
            # doc.save(output_file)
            
            # Create success message
            message = f"Rapor kaydedildi: {output_file}\n" + \
                    f"Güncellenen toplam satır: {total_updated_count}\n" + \
                    f"Güncellenen tablo sayısı: {updated_tables_count}"
            
            if lot_details_file:
                message += f"\nLot detayları: {lot_details_file}"
            
            if lot_details_pdf:
                message += f"\nLot detayları PDF: {lot_details_pdf}"
            
            if nonconforming_count:
                message += f"\nUygunsuz ölçüm sayısı: {nonconforming_count}"
            
            logger.info(f"Report updated/created: {output_file}")
            return True, output_file, message
            
        except Exception as e:
            logger.error(f"Error creating report: {str(e)}")
            return False, None, f"Error: {str(e)}"

    def _find_header_row(self, table):
        """
        Find the header row in the table
        
        Args:
            table: Word table object
            
        Returns:
            int: Header row index or -1 if not found
        """
        for i, row in enumerate(table.rows):
            # Get text from all cells
            row_text = " ".join([cell.text.strip().upper() for cell in row.cells])
            
            # Skip RECORDING SHEET rows
            if any(pattern in row_text for pattern in ["RECORDIG SHEET", "RECORDING SHEET", "REC ORDIG SHEET"]):
                continue
            
            # Look for header indicators
            if ("ITEM" in row_text and "NO" in row_text) and "DIMENSION" in row_text and "ACTUAL" in row_text:
                return i
        
        # If not found, check first few rows for ACTUAL
        for i in range(min(3, len(table.rows))):
            header_cells = [cell.text.strip().upper() for cell in table.rows[i].cells]
            if any("ACTUAL" in cell for cell in header_cells):
                return i
        
        return -1
    
    def _get_column_indices_for_table(self, table, header_row_idx):
        """
        Find important column indices in the table header row
        
        Args:
            table: Word table object
            header_row_idx: Index of the header row
            
        Returns:
            dict: Dictionary of column indices
        """
        indices = {
            'actual': -1,
            'item_no': -1,
            'dimension': -1,
            'badge': -1,
            'tooling': -1,
            'insp_level': -1,
            'bp_zone': -1  # Added for nonconforming item format
        }
        
        # Check if header row exists
        if header_row_idx >= 0 and header_row_idx < len(table.rows):
            header_row = table.rows[header_row_idx]
            
            # Find indices in current table's header
            for idx, cell in enumerate(header_row.cells):
                header_upper = cell.text.strip().upper()
                
                if "ACTUAL" in header_upper:
                    indices['actual'] = idx
                if "ITEM NO" in header_upper or "ITEMNO" in header_upper:
                    indices['item_no'] = idx
                if "DIMENSION" in header_upper or "DIM" in header_upper:
                    indices['dimension'] = idx
                if "BADGE" in header_upper:
                    indices['badge'] = idx
                if "TOOLING" in header_upper:
                    indices['tooling'] = idx
                if "INSP. LEVEL" in header_upper or "INSPLEVEL" in header_upper:
                    indices['insp_level'] = idx
                if "B/P ZONE" in header_upper or "BP ZONE" in header_upper or "B/P" in header_upper:
                    indices['bp_zone'] = idx
        
        return indices
    
    def _update_actual_values(self, table, header_row_idx, actual_values, col_indices):
        """
        Update ACTUAL values in the table
        
        Args:
            table: Word table object
            header_row_idx (int): Header row index
            actual_values (dict): Dictionary of actual values
            col_indices (dict): Dictionary of column indices
            
        Returns:
            int: Number of updated rows
        """
        updated_count = 0
        
        # Get column indices from indices dict
        actual_col_idx = col_indices['actual']
        item_no_col_idx = col_indices['item_no']
        dimension_col_idx = col_indices['dimension']
        
        # Make sure we have the required columns
        if actual_col_idx == -1 or item_no_col_idx == -1 or dimension_col_idx == -1:
            logger.error("Required columns not found")
            return 0
        
        # Update values
        for i, row in enumerate(table.rows):
            # Skip header row and anything before it
            if i <= header_row_idx:
                continue
            
            # Check column indices are within range
            if (item_no_col_idx >= len(row.cells) or 
                dimension_col_idx >= len(row.cells) or 
                actual_col_idx >= len(row.cells)):
                continue
            
            # Get cell values
            item_no = row.cells[item_no_col_idx].text.strip()
            dimension = row.cells[dimension_col_idx].text.strip()
            
            # Skip if item_no or dimension is empty
            if not item_no or not dimension:
                continue
            
            # Create key
            key = f"{dimension}_{item_no}"
            
            # Check if we have a value for this key
            if key in actual_values and actual_values[key]:
                actual_value = str(actual_values[key])
                
                # Update the cell
                row.cells[actual_col_idx].text = actual_value
                updated_count += 1
                logger.debug(f"Updated {key} with {actual_value}")
        
        return updated_count


class TableViewer(ctk.CTk):
    """
    Main application class for viewing and editing table data
    """
    def __init__(self):
        super().__init__()
        
        # Initialize app components
        self.title("Tam Tablo Görüntüleyici - ACTUAL ile EntryBox")
        self.geometry("1200x900")
        
        # Initialize managers
        self.project_manager = ProjectManager()
        self.table_reader = TableReader()
        self.lot_detail_manager = LotDetailManager(self)
        self.report_generator = ReportGenerator(self.project_manager)
        
        # Initialize data
        self.actual_values = {}  # Store ACTUAL values
        self.widgets = {}  # Store UI widgets
        self.checkbox_states = {}  # Store checkbox states
        self.checkbox_widgets = {}  # Store checkbox widgets
        self.nonconforming_items = {}  # Store nonconforming item messages
        self.row_data = {}  # Store row data for each key
        self.selected_table_idx = 0
        self.selected_file_path = None
        self.headers = []
        self.col_indices = {}
        
        # Set up UI
        self._setup_ui()
 
    
    def _setup_ui(self):
        """Set up the user interface with TabView"""
        # Main frame
        self.main_frame = ctk.CTkFrame(self, corner_radius=10)
        self.main_frame.pack(padx=30, pady=30, fill="both", expand=True)
        
        # Main title
        self.title_label = ctk.CTkLabel(
            self.main_frame, 
            text="Ölçüm Rapor Sistemi", 
            font=("Helvetica", 24, "bold")
        )
        self.title_label.pack(pady=(0, 20))
        
        # Create TabView
        self.tabview = ctk.CTkTabview(self.main_frame)
        self.tabview.pack(fill="both", expand=True, padx=10, pady=10)
        
        # Add tabs
        self.tab_project = self.tabview.add("Proje Bilgileri")
        self.tab_table = self.tabview.add("Tablo Görüntüleme")
        self.tab_report = self.tabview.add("Rapor Oluşturma")
        
        # Set default tab
        self.tabview.set("Proje Bilgileri")
        
        # Setup each tab
        self._setup_project_tab()
        self._setup_table_tab()
        self._setup_report_tab()


    def _setup_project_tab(self):
        """Set up the project information tab"""
        # Project info frame
        self.info_frame = ctk.CTkFrame(self.tab_project, corner_radius=8)
        self.info_frame.pack(fill="both", expand=True, padx=20, pady=20)
        
        # Title
        ctk.CTkLabel(
            self.info_frame, 
            text="Proje Bilgileri", 
            font=("Helvetica", 20, "bold")
        ).pack(pady=(20, 30))
        
        # Create grid frame for inputs
        input_frame = ctk.CTkFrame(self.info_frame, corner_radius=0)
        input_frame.pack(fill="x", padx=40, pady=10)
        
        # Project type
        ctk.CTkLabel(input_frame, text="Proje Tipi:", font=("Helvetica", 14)).grid(
            row=0, column=0, padx=10, pady=15, sticky="e"
        )
        self.project_type = ctk.CTkComboBox(
            input_frame, 
            values=["Tip A", "Tip B", "Tip C"], 
            font=("Helvetica", 14), 
            width=250
        )
        self.project_type.grid(row=0, column=1, padx=10, pady=15, sticky="w")
        
        # Part number
        ctk.CTkLabel(input_frame, text="Parça Numarası:", font=("Helvetica", 14)).grid(
            row=1, column=0, padx=10, pady=15, sticky="e"
        )
        self.part_number = ctk.CTkEntry(input_frame, font=("Helvetica", 14), width=250)
        self.part_number.grid(row=1, column=1, padx=10, pady=15, sticky="w")
        
        # Operation number
        ctk.CTkLabel(input_frame, text="Operasyon No:", font=("Helvetica", 14)).grid(
            row=2, column=0, padx=10, pady=15, sticky="e"
        )
        self.operation_number = ctk.CTkEntry(input_frame, font=("Helvetica", 14), width=250)
        self.operation_number.grid(row=2, column=1, padx=10, pady=15, sticky="w")
        
        # Serial number
        ctk.CTkLabel(input_frame, text="Seri No:", font=("Helvetica", 14)).grid(
            row=3, column=0, padx=10, pady=15, sticky="e"
        )
        self.serial_number = ctk.CTkEntry(input_frame, font=("Helvetica", 14), width=250)
        self.serial_number.grid(row=3, column=1, padx=10, pady=15, sticky="w")
        
        # Continue measurement checkbox
        self.continue_measurement = ctk.CTkCheckBox(
            input_frame, 
            text="Yarım Kalan Ölçüme Devam Et", 
            font=("Helvetica", 14)
        )
        self.continue_measurement.grid(row=4, column=0, columnspan=2, padx=10, pady=20)

        # Save button
        self.save_info_button = ctk.CTkButton(
            input_frame, 
            text="Bilgileri Kaydet", 
            font=("Helvetica", 16),
            width=200, 
            height=40, 
            corner_radius=8,
            command=self.save_project_info
        )
        self.save_info_button.grid(row=5, column=0, columnspan=2, padx=10, pady=20)
        
        # Configure grid columns
        input_frame.grid_columnconfigure(0, weight=1)
        input_frame.grid_columnconfigure(1, weight=1)
        
        # File selection section - MOVED FROM TABLE TAB TO PROJECT TAB
        file_frame = ctk.CTkFrame(self.info_frame, corner_radius=0)
        file_frame.pack(fill="x", padx=40, pady=10)
        
        # File selection title
        ctk.CTkLabel(
            file_frame, 
            text="Dosya Seçimi", 
            font=("Helvetica", 16, "bold")
        ).pack(pady=(10, 15))
        
        # File selection button
        self.select_button = ctk.CTkButton(
            file_frame, 
            text="Word Dosyası Seç", 
            font=("Helvetica", 16),
            width=200, 
            height=40, 
            corner_radius=8,
            command=self.select_file
        )
        self.select_button.pack(pady=15)
        
        # File path display
        self.file_path_var = tk.StringVar(value="Henüz dosya seçilmedi")
        ctk.CTkLabel(
            file_frame, 
            textvariable=self.file_path_var,
            font=("Helvetica", 12, "italic")
        ).pack(pady=(0, 15))
        
        # Button frame
        button_frame = ctk.CTkFrame(self.info_frame, corner_radius=0)
        button_frame.pack(fill="x", padx=40, pady=(20, 40))


    def _setup_table_tab(self):
        """Set up the table viewing tab"""
        # Table frame
        self.table_content_frame = ctk.CTkFrame(self.tab_table, corner_radius=8)
        self.table_content_frame.pack(fill="both", expand=True, padx=20, pady=20)
        
        # Table selector frame
        self.table_selector_frame = ctk.CTkFrame(self.table_content_frame, corner_radius=0)
        self.table_selector_frame.pack(fill="x", padx=40, pady=10)
        
        # Table selector label
        ctk.CTkLabel(
            self.table_selector_frame,
            text="Tablo Seçiniz:",
            font=("Helvetica", 14)
        ).pack(side="left", padx=(0, 10))
        
        # Table selector
        self.table_selector = ctk.CTkComboBox(
            self.table_selector_frame, 
            values=[], 
            font=("Helvetica", 14),
            width=200, 
            height=30, 
            corner_radius=6,
            command=self.on_table_select
        )
        self.table_selector.pack(side="left", padx=10)
        
        # Hide initially
        self.table_selector_frame.pack_forget()
        
        # Create a frame for the table content
        self.table_display_frame = ctk.CTkFrame(self.table_content_frame, corner_radius=8)
        self.table_display_frame.pack(fill="both", expand=True, padx=20, pady=20)
        
        # Initial message
        self.table_initial_message = ctk.CTkLabel(
            self.table_display_frame,
            text="Lütfen önce bir Word dosyası seçiniz.",
            font=("Helvetica", 16)
        )
        self.table_initial_message.pack(pady=50)


    def _setup_report_tab(self):
        """Set up the report generation tab"""
        # Report frame
        self.report_frame = ctk.CTkFrame(self.tab_report, corner_radius=8)
        self.report_frame.pack(fill="both", expand=True, padx=20, pady=20)
        
        # Title
        ctk.CTkLabel(
            self.report_frame, 
            text="Rapor Oluşturma", 
            font=("Helvetica", 20, "bold")
        ).pack(pady=(20, 30))
        
        # Report info frame
        info_display = ctk.CTkFrame(self.report_frame, corner_radius=6)
        info_display.pack(fill="x", padx=40, pady=10)
        
        # Project info section
        ctk.CTkLabel(
            info_display,
            text="Proje Bilgileri",
            font=("Helvetica", 16, "bold")
        ).pack(anchor="w", padx=20, pady=(20, 10))
        
        # Project info display
        self.project_info_text = ctk.CTkTextbox(info_display, height=100, corner_radius=6)
        self.project_info_text.pack(fill="x", padx=20, pady=10)
        self.project_info_text.configure(state="disabled")  # Read-only
        
        # File info section
        ctk.CTkLabel(
            info_display,
            text="Dosya Bilgileri",
            font=("Helvetica", 16, "bold")
        ).pack(anchor="w", padx=20, pady=(20, 10))
        
        # File info display
        self.file_info_text = ctk.CTkTextbox(info_display, height=60, corner_radius=6)
        self.file_info_text.pack(fill="x", padx=20, pady=10)
        self.file_info_text.configure(state="disabled")  # Read-only
        
        # Summary section
        ctk.CTkLabel(
            info_display,
            text="Özet Bilgiler",
            font=("Helvetica", 16, "bold")
        ).pack(anchor="w", padx=20, pady=(20, 10))
        
        # Summary display
        self.summary_text = ctk.CTkTextbox(info_display, height=80, corner_radius=6)
        self.summary_text.pack(fill="x", padx=20, pady=10)
        self.summary_text.configure(state="disabled")  # Read-only
        
        # Buttons frame
        button_frame = ctk.CTkFrame(self.report_frame, corner_radius=0)
        button_frame.pack(fill="x", padx=40, pady=(20, 40))
        
        # Generate report button
        self.generate_report_button = ctk.CTkButton(
            button_frame, 
            text="Rapor Oluştur", 
            font=("Helvetica", 16),
            width=200, 
            height=40, 
            corner_radius=8,
            command=self.create_report
        )
        self.generate_report_button.pack(pady=10)
        
        # Go back button
        ctk.CTkButton(
            button_frame, 
            text="Tablo Görüntülemeye Dön", 
            font=("Helvetica", 14),
            width=200, 
            height=35, 
            corner_radius=8,
            command=lambda: self.tabview.set("Tablo Görüntüleme")
        ).pack(pady=10)


    def save_project_info(self):
        """Save project information and create folder structure"""
        # Get values
        project_type = self.project_type.get()
        part_number = self.part_number.get()
        operation_number = self.operation_number.get()
        serial_number = self.serial_number.get()
        continue_measurement = self.continue_measurement.get()
        
        # If continuing measurement, check if folder exists and load lot details
        if continue_measurement:
            folder_exists, folder_path = self.project_manager.initialize_from_existing(
                project_type, 
                part_number, 
                operation_number, 
                serial_number
            )
            
            if folder_exists:
                # Try to load existing lot details
                self.lot_detail_manager.load_lot_details_from_text(folder_path)
                
                # Update report tab
                self.update_report_info()
                
                # Show success message
                messagebox.showinfo("Bilgi", f"Varolan proje bilgileri yüklendi!\nKlasör: {folder_path}")
                
                # No need to create new structure, return
                return
            else:
                # Folder doesn't exist, continue with normal creation
                messagebox.showinfo("Bilgi", "Belirtilen proje klasörü bulunamadı. Yeni klasör oluşturulacak.")
        
        # Create project structure
        success, folder_path, error_message = self.project_manager.create_project_structure(
            project_type, 
            part_number, 
            operation_number, 
            serial_number, 
            continue_measurement
        )
        
        if not success:
            messagebox.showwarning("Uyarı", error_message or "LÜTFEN TÜM BİLGİLERİ DOLDURUNUZ")
            return
        
        # Open the folder
        self.project_manager.open_folder(folder_path)
        
        # Update report tab
        self.update_report_info()
        
        # Show success message
        messagebox.showinfo("Bilgi", f"Proje bilgileri kaydedildi!\nKlasör: {folder_path}")
    

    def select_file(self):
        """Select and load a Word document"""
        # Check if we need project info first
        if not self.continue_measurement.get() and not self.project_manager.serial_folder:
            messagebox.showwarning("Uyarı", "Lütfen önce proje bilgilerini kaydediniz!")
            self.tabview.set("Proje Bilgileri")
            return
        
        # Open file dialog
        file_path = filedialog.askopenfilename(filetypes=[("Word dosyaları", "*.docx")])
        if not file_path:
            return
        
        try:
            # Store file path
            self.selected_file_path = file_path
            self.file_path_var.set(f"Seçilen dosya: {os.path.basename(file_path)}")
            
            # Read tables
            self.table_reader.read_word_tables(file_path)
            
            # Update UI
            self.update_table_selector()
            
            # Show table selector
            self.table_selector_frame.pack(fill="x", padx=40, pady=10)
            
            # Clear the initial message
            self.table_initial_message.pack_forget()
            
            # Load lot details if available (for continuing work)
            if self.project_manager.serial_folder:
                self.lot_detail_manager.load_lot_details_from_text(self.project_manager.serial_folder)
                logger.info("Checked for existing lot details")
            
            # Show first table
            self.show_table()
            
            # Update report tab
            self.update_report_info()
            
            # Auto switch to next tab
            self.tabview.set("Tablo Görüntüleme")
            
        except Exception as e:
            logger.error(f"Error selecting file: {str(e)}")
            messagebox.showerror("Hata", f"Dosya okuma hatası: {str(e)}")
    

    def update_table_selector(self):
        """Update the table selector dropdown with available tables"""
        if self.table_reader.table_data:
            table_names = [f"Tablo {i+1}" for i in range(len(self.table_reader.table_data))]
            self.table_selector.configure(values=table_names)
            self.table_selector.set(table_names[0])
    

    def on_table_select(self, choice):
        """Handle table selection from dropdown"""
        try:
            table_idx = self.table_selector.get().split()[1]
            self.selected_table_idx = int(table_idx) - 1
            self.show_table()
            self.update_report_info()
        except Exception as e:
            logger.error(f"Error selecting table: {str(e)}")


    def show_table(self):
        """Display the selected table"""
        # Clear previous table content
        for widget in self.table_display_frame.winfo_children():
            widget.destroy()
        
        # Check if we have data
        if not self.table_reader.table_data or self.selected_table_idx >= len(self.table_reader.table_data):
            ctk.CTkLabel(self.table_display_frame, text="Tablo verisi bulunamadı.", font=("Helvetica", 16)).pack(pady=50)
            return
        
        # Get selected table
        table = self.table_reader.table_data[self.selected_table_idx]
        if not table or not table[0]:
            return
        
        # Get headers
        self.headers = table[0].copy()
        
        # Add "Lot Check" header if not present
        if "Lot Check" not in self.headers:
            self.headers.append("Lot Check")
        
        # Get column indices
        self.col_indices = self.table_reader.get_column_indices(self.headers)
        
        # Check required columns
        if self.col_indices['actual'] == -1:
            ctk.CTkLabel(self.table_display_frame, text="Tabloda ACTUAL sütunu bulunamadı.", font=("Helvetica", 16)).pack(pady=50)
            return
        if self.col_indices['dimension'] == -1 or self.col_indices['item_no'] == -1:
            ctk.CTkLabel(self.table_display_frame, text="Tabloda DIMENSION veya ITEM NO sütunu bulunamadı.", font=("Helvetica", 16)).pack(pady=50)
            return
        
        # Create scrollable frame for table
        self.table_frame = ctk.CTkScrollableFrame(self.table_display_frame, corner_radius=10, orientation="horizontal")
        self.table_frame.pack(pady=20, padx=20, fill="both", expand=True)
        
        # Calculate column widths
        self._calculate_column_widths(table)
        
        # Render table
        self._render_table(table)
        
        # Add buttons
        self._add_action_buttons()


    def _calculate_column_widths(self, table):
        """Calculate optimal column widths based on content"""
        # Initialize with length of headers
        self.col_widths = [len(str(header)) for header in self.headers]
        
        # Update with data
        for row in table:
            for j, cell in enumerate(row):
                if j < len(self.col_widths):
                    self.col_widths[j] = max(self.col_widths[j], len(str(cell)))
        
        # Convert to pixel widths
        window_width = self.winfo_screenwidth()
        num_cols = len(self.headers)
        base_width = max(120, min(window_width // num_cols, max(self.col_widths) * 12))
        self.col_widths = [max(120, min(base_width, w * 12)) for w in self.col_widths]


    def _render_table(self, table):
            """Render the table with headers and data"""
            # First render headers
            header_row = 0
            # Add a new header for the checkbox column
            checkbox_header = ctk.CTkLabel(
                self.table_frame, 
                text="Seç", 
                font=("Helvetica", 14, "bold"),
                width=60, 
                anchor="center", 
                corner_radius=6,
                bg_color="transparent"
            )
            checkbox_header.grid(row=header_row, column=0, padx=5, pady=5, sticky="nsew")
            
            for j, header in enumerate(self.headers):
                # Skip BADGE column if present
                if j == self.col_indices['badge']:
                    continue
                
                # Create header label
                label = ctk.CTkLabel(
                    self.table_frame, 
                    text=header, 
                    font=("Helvetica", 14, "bold"),
                    width=self.col_widths[j] if j < len(self.col_widths) else 120, 
                    anchor="center", 
                    corner_radius=6,
                    bg_color="transparent"
                )
                
                # Adjust grid column if we're skipping BADGE, and add 1 for checkbox column
                grid_col = j + 1 if j < self.col_indices['badge'] else j
                if j == len(self.headers) - 1:  # "Lot Check" column
                    grid_col = len(self.headers) - (1 if self.col_indices['badge'] != -1 else 0)
                
                label.grid(row=header_row, column=grid_col, padx=5, pady=5, sticky="nsew")
            
            # Then render data rows
            for i, row in enumerate(table[1:], 1):  # Skip header row
                # Add checkbox at the beginning of each row
                if (self.col_indices['dimension'] < len(row) and 
                    self.col_indices['item_no'] < len(row)):
                    
                    dimension = row[self.col_indices['dimension']]
                    item_no = row[self.col_indices['item_no']]
                    key = f"{dimension}_{item_no}"
                    
                    # Store row data for reference
                    row_data = {
                        'item_no': item_no,
                        'dimension': dimension,
                        'actual': row[self.col_indices['actual']] if self.col_indices['actual'] < len(row) else "",
                        'bp_zone': row[self.col_indices.get('bp_zone', -1)] if 'bp_zone' in self.col_indices and self.col_indices['bp_zone'] < len(row) else ""
                    }
                    self.row_data[key] = row_data
                    
                    # Create a checkbox variable
                    checkbox_var = tk.IntVar(value=self.checkbox_states.get(key, 0))
                    
                    # Create checkbox
                    checkbox = ctk.CTkCheckBox(
                        self.table_frame,
                        text="",
                        variable=checkbox_var,
                        width=40,
                        height=20,
                        corner_radius=3,
                        command=lambda k=key, v=checkbox_var: self.on_checkbox_toggle(k, v)
                    )
                    checkbox.grid(row=i, column=0, padx=5, pady=5, sticky="nsew")
                    
                    # Store checkbox widget and state
                    self.checkbox_widgets[key] = checkbox
                    self.checkbox_states[key] = checkbox_var.get()
                
                for j, cell in enumerate(row):
                    # Skip BADGE column
                    if j == self.col_indices['badge']:
                        continue
                    
                    # Adjust grid column (add 1 for checkbox column)
                    grid_col = j + 1 if j < self.col_indices['badge'] else j
                    
                    # Ensure j is within valid range for row
                    if j >= len(row):
                        continue
                    
                    if j == self.col_indices['actual']:  # ACTUAL column as entry
                        # Create key if dimension and item_no indices are valid
                        if (self.col_indices['dimension'] < len(row) and 
                            self.col_indices['item_no'] < len(row)):
                            
                            dimension = row[self.col_indices['dimension']]
                            item_no = row[self.col_indices['item_no']]
                            key = f"{dimension}_{item_no}"
                            
                            # Create entry widget
                            entry = ctk.CTkEntry(
                                self.table_frame, 
                                width=self.col_widths[j] if j < len(self.col_widths) else 120, 
                                height=30, 
                                corner_radius=6,
                                font=("Helvetica", 14)
                            )
                            
                            # Set initial value - prioritize saved actual_values over the cell content
                            initial_value = self.actual_values.get(key, "") or cell
                            entry.insert(0, initial_value)
                            
                            # Bind events
                            entry.bind("<Return>", lambda e, k=key, w=entry: self.update_actual_value(k, w.get()))
                            entry.bind("<FocusOut>", lambda e, k=key, w=entry: self.update_actual_value(k, w.get()))
                            
                            # Store widget reference
                            self.widgets[key] = entry
                            
                            # Add to grid
                            entry.grid(row=i, column=grid_col, padx=5, pady=5, sticky="nsew")
                        else:
                            # Fallback if indices are invalid
                            label = ctk.CTkLabel(
                                self.table_frame, 
                                text=cell or "", 
                                font=("Helvetica", 14),
                                width=self.col_widths[j] if j < len(self.col_widths) else 120, 
                                anchor="center", 
                                corner_radius=6,
                                bg_color="transparent"
                            )
                            label.grid(row=i, column=grid_col, padx=5, pady=5, sticky="nsew")
                    else:
                        # Regular cell as label
                        label = ctk.CTkLabel(
                            self.table_frame, 
                            text=cell or "", 
                            font=("Helvetica", 14),
                            width=self.col_widths[j] if j < len(self.col_widths) else 120, 
                            anchor="center", 
                            corner_radius=6,
                            bg_color="transparent"
                        )
                        label.grid(row=i, column=grid_col, padx=5, pady=5, sticky="nsew")
                
                # Add "Detail" button for Lot Check column if dimension and item_no indices are valid
                if (self.col_indices['dimension'] < len(row) and 
                    self.col_indices['item_no'] < len(row)):
                    
                    dimension = row[self.col_indices['dimension']]
                    item_no = row[self.col_indices['item_no']]
                    
                    detail_button = ctk.CTkButton(
                        self.table_frame, 
                        text="Detail", 
                        font=("Helvetica", 12),
                        width=self.col_widths[-1] - 20 if len(self.col_widths) > 0 else 100, 
                        height=30, 
                        corner_radius=6,
                        command=lambda r=i, item=item_no, dim=dimension: self.show_lot_detail(r, item, dim)
                    )
                    
                    # Calculate grid column (offset by 1 for checkbox column)
                    detail_grid_col = len(self.headers) - (1 if self.col_indices['badge'] != -1 else 0)
                    
                    detail_button.grid(row=i, column=detail_grid_col, padx=5, pady=5, sticky="nsew")
            
            # Configure grid
            visible_cols = len(self.headers) - (1 if self.col_indices['badge'] != -1 else 0) + 1  # +1 for checkbox col
            self.table_frame.grid_columnconfigure(0, weight=0, minsize=60)  # Checkbox column
            for j in range(1, visible_cols):
                self.table_frame.grid_columnconfigure(j, weight=1, minsize=120)
            
            for i in range(len(table)):
                self.table_frame.grid_rowconfigure(i, weight=1)


    def on_checkbox_toggle(self, key, checkbox_var):
        """Handle checkbox toggle events"""
        self.checkbox_states[key] = checkbox_var.get()
        
        # Generate nonconforming message if checkbox is checked
        if self.checkbox_states[key] == 1 and key in self.row_data:
            row = self.row_data[key]
            item_no = row.get('item_no', '')
            dimension = row.get('dimension', '')
            bp_zone = row.get('bp_zone', '')
            
            # Get actual value - first check widgets for most current value
            actual = ""
            if key in self.widgets:
                actual = self.widgets[key].get()
                
            # If not found in widgets, try stored values
            if not actual:
                actual = self.actual_values.get(key, '') or row.get('actual', '')
            
            # Format: "[ITEM NO] DIMENSION (B/P ZONE) checks ACTUAL."
            bp_zone_text = f" ({bp_zone})" if bp_zone else ""
            nonconforming_message = f"[{item_no}] {dimension}{bp_zone_text} checks {actual}."
            
            # Store the message
            self.nonconforming_items[key] = nonconforming_message
            logger.info(f"Nonconforming item: {nonconforming_message}")
        elif key in self.nonconforming_items and self.checkbox_states[key] == 0:
            # Remove from nonconforming items if unchecked
            del self.nonconforming_items[key]
            logger.info(f"Removed item {key} from nonconforming items")
                
        logger.debug(f"Checkbox toggled for {key}: {self.checkbox_states[key]}")

    def _add_action_buttons(self):
        """Add save and report buttons"""
        button_frame = ctk.CTkFrame(self.table_frame)
        button_frame.grid(
            row=len(self.table_reader.table_data[self.selected_table_idx]), 
            column=0, 
            columnspan=len(self.headers) - (1 if self.col_indices['badge'] != -1 else 0), 
            pady=10, 
            padx=5, 
            sticky="w"
        )
        
        save_button = ctk.CTkButton(
            button_frame, 
            text="Değerleri Kaydet", 
            font=("Helvetica", 14),
            width=150, 
            height=35, 
            corner_radius=6,
            command=self.save_values
        )
        save_button.pack(side="left", padx=10)
        
        go_to_report_button = ctk.CTkButton(
            button_frame, 
            text="Rapor Sekmesine Geç", 
            font=("Helvetica", 14),
            width=170, 
            height=35, 
            corner_radius=6,
            command=lambda: [self.update_report_info(), self.tabview.set("Rapor Oluşturma")]
        )
        go_to_report_button.pack(side="left", padx=10)
    

    def show_lot_detail(self, row_idx, item_no, dimension):
        """Show lot detail dialog"""
        # Get actual value
        key = f"{dimension}_{item_no}"
        actual_value = self.actual_values.get(key, "")
        
        # Show dialog
        self.lot_detail_manager.show_lot_detail_dialog(
            row_idx, 
            item_no, 
            dimension, 
            actual_value, 
            self.widgets
        )
        
        # Update report tab after closing the dialog
        self.update_report_info()
    

    def update_actual_value(self, key, value):
        """Update ACTUAL value"""
        self.actual_values[key] = value
        logger.debug(f"Updated actual value for {key}: {value}")
        
        # Update report tab
        self.update_report_info()
    

    def save_values(self):
        """Save all ACTUAL values"""
        for key, value in self.actual_values.items():
            logger.debug(f"Saved: {key} -> {value}")
        
        # Update report tab
        self.update_report_info()
        
        messagebox.showinfo("Başarılı", "Değerler başarıyla kaydedildi!")
    

    def update_report_info(self):
            """Update the report tab with current information"""
            # Update project info
            self.project_info_text.configure(state="normal")
            self.project_info_text.delete("1.0", "end")
            
            if self.project_manager.project_info:
                for key, value in self.project_manager.project_info.items():
                    self.project_info_text.insert("end", f"{key}: {value}\n")
            else:
                self.project_info_text.insert("end", "Proje bilgileri henüz kaydedilmedi.\n")
            
            self.project_info_text.configure(state="disabled")
            
            # Update file info
            self.file_info_text.configure(state="normal")
            self.file_info_text.delete("1.0", "end")
            
            if self.selected_file_path:
                self.file_info_text.insert("end", f"Dosya: {os.path.basename(self.selected_file_path)}\n")
                self.file_info_text.insert("end", f"Tablo Sayısı: {len(self.table_reader.table_data)}\n")
                self.file_info_text.insert("end", f"Seçili Tablo: {self.selected_table_idx + 1}\n")
            else:
                self.file_info_text.insert("end", "Henüz dosya seçilmedi.\n")
            
            self.file_info_text.configure(state="disabled")
            
            # Update summary
            self.summary_text.configure(state="normal")
            self.summary_text.delete("1.0", "end")
            
            actual_count = len(self.actual_values)
            lot_count = len(self.lot_detail_manager.part_quantities)
            
            self.summary_text.insert("end", f"Girilen ACTUAL değer sayısı: {actual_count}\n")
            self.summary_text.insert("end", f"Lot detay sayısı: {lot_count}\n")
            
            if not self.project_manager.serial_folder:
                self.summary_text.insert("end", "UYARI: Proje bilgileri kaydedilmeden rapor oluşturulamaz.\n")
            elif not self.selected_file_path:
                self.summary_text.insert("end", "UYARI: Dosya seçilmeden rapor oluşturulamaz.\n")
            elif actual_count == 0:
                self.summary_text.insert("end", "UYARI: Değer girilmeden rapor oluşturulabilir ancak tabloda değişiklik olmaz.\n")
            
            self.summary_text.configure(state="disabled")


    def create_report(self):
        """Create a report document and export nonconforming items to Excel"""
        # Validate project info
        if not self.project_manager.serial_folder:
            messagebox.showerror("Hata", "Önce proje bilgilerini kaydetmelisiniz!")
            self.tabview.set("Proje Bilgileri")
            return
        
        # Validate file selection
        if not self.selected_file_path:
            messagebox.showerror("Hata", "Önce bir dosya seçmelisiniz!")
            self.tabview.set("Tablo Görüntüleme")
            return
        
        # Confirm report generation
        actual_count = len(self.actual_values)
        if actual_count == 0:
            proceed = messagebox.askyesno(
                "Uyarı", 
                "Hiç ACTUAL değeri girilmemiş. Rapor oluşturmaya devam etmek istiyor musunuz?"
            )
            if not proceed:
                return
        
        # Add nonconforming items to project info for report
        nonconforming_count = len(self.nonconforming_items)
        if nonconforming_count > 0:
            self.project_manager.project_info['UYGUNSUZ_OLCUM_SAYISI'] = str(nonconforming_count)
            
            # Create a formatted string of nonconforming items for the report
            nonconforming_text = ""
            for idx, message in enumerate(self.nonconforming_items.values(), 1):
                nonconforming_text += f"{idx}. {message}\n"
            
            self.project_manager.project_info['UYGUNSUZ_OLCUMLER'] = nonconforming_text
        
        try:
            # Generate report
            success, file_path, message = self.report_generator.create_report(
                self.selected_file_path,
                self.selected_table_idx,
                self.actual_values,
                self.col_indices,
                self.lot_detail_manager
            )
            
            # Export nonconforming items to Excel if any exist
            excel_file = None
            if nonconforming_count > 0:
                excel_file = self.export_nonconforming_to_excel()
                
                if excel_file:
                    message += f"\nUygunsuz ölçümler Excel dosyası: {os.path.basename(excel_file)}"
            
            # Show result
            if success:
                # Add nonconforming items to the success message
                if nonconforming_count > 0:
                    nonconforming_text = "\n\nUYGUNSUZ ÖLÇÜMLER:\n"
                    for idx, msg in enumerate(self.nonconforming_items.values(), 1):
                        nonconforming_text += f"{idx}. {msg}\n"
                    message += nonconforming_text
                
                messagebox.showinfo("Başarılı", message)
                
                # Ask if user wants to open the files
                # if excel_file and nonconforming_count > 0:
                #     open_excel = messagebox.askyesno(
                #         "Bilgi", 
                #         "Uygunsuz ölçümler Excel dosyası oluşturuldu. Açmak ister misiniz?"
                #     )
                #     if open_excel:
                #         try:
                #             if os.name == 'nt':  # Windows
                #                 os.startfile(excel_file)
                #             elif os.name == 'posix':  # macOS, Linux
                #                 if sys.platform == 'darwin':  # macOS
                #                     subprocess.call(['open', excel_file])
                #                 else:  # Linux
                #                     subprocess.call(['xdg-open', excel_file])
                #         except Exception as e:
                #             logger.error(f"Error opening Excel file: {str(e)}")
                #             messagebox.showerror("Hata", f"Excel dosyası açılırken hata oluştu: {str(e)}")
                
                open_file = messagebox.askyesno("Bilgi", "Oluşturulan raporu açmak ister misiniz?")
                if open_file:
                    try:
                        if os.name == 'nt':  # Windows
                            os.startfile(file_path)
                        elif os.name == 'posix':  # macOS, Linux
                            if sys.platform == 'darwin':  # macOS
                                subprocess.call(['open', file_path])
                            else:  # Linux
                                subprocess.call(['xdg-open', file_path])
                    except Exception as e:
                        logger.error(f"Error opening file: {str(e)}")
                        messagebox.showerror("Hata", f"Dosya açılırken hata oluştu: {str(e)}")
            else:
                messagebox.showerror("Hata", message)
        except Exception as e:
            logger.error(f"Error creating report: {str(e)}")
            messagebox.showerror("Hata", f"Rapor oluşturulurken hata oluştu: {str(e)}")


    def export_nonconforming_to_excel(self):
        """
        Export nonconforming items to a simple Excel file.
        Creates a file with just the nonconforming item messages in the specified format.
        
        Returns:
            str: Path to the created Excel file or None if failed
        """
        try:
            # Check if we have any nonconforming items
            if not self.nonconforming_items:
                logger.info("No nonconforming items to export to Excel")
                return None
            
            # Import openpyxl
            try:
                import openpyxl
                from openpyxl.styles import Font, Alignment
            except ImportError:
                logger.error("Openpyxl library not found. Required for Excel export.")
                return None
            
            # Check if project folder is created
            if not self.project_manager.serial_folder:
                logger.error("Project folder not created. Cannot export to Excel.")
                return None
            
            # Create Excel file path
            timestamp = datetime.datetime.now().strftime('%Y%m%d_%H%M%S')
            excel_file = os.path.join(self.project_manager.serial_folder, f"Uygunsuz_Olcumler_{timestamp}.xlsx")
            
            # Create a new workbook and select the active sheet
            wb = openpyxl.Workbook()
            ws = wb.active
            ws.title = "Uygunsuz Ölçümler"
            
            # Add a simple header
            ws['A1'] = "UYGUNSUZ ÖLÇÜMLER"
            ws['A1'].font = Font(bold=True, size=14)
            
            # Add the nonconforming item messages
            row = 3
            for message in self.nonconforming_items.values():
                ws[f'A{row}'] = message
                row += 1
            
            # Auto-adjust column width
            max_length = 0
            for cell in ws['A']:
                if cell.value:
                    try:
                        cell_length = len(str(cell.value))
                        if cell_length > max_length:
                            max_length = cell_length
                    except:
                        pass
            ws.column_dimensions['A'].width = max_length + 5
            
            # Save the workbook
            wb.save(excel_file)
            logger.info(f"Created simplified Excel file with nonconforming items: {excel_file}")
            
            return excel_file
            
        except Exception as e:
            logger.error(f"Error exporting to Excel: {str(e)}")
            return None

def main():
    app = TableViewer()
    app.mainloop()

if __name__ == "__main__":
    main()